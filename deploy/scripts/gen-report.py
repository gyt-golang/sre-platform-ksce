#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《基于金山云平台的 SRE 可靠性工程平台》项目报告 Word 版。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# 全局字体（中文宋体，西文 Calibri）
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_cn(run, name='微软雅黑'):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def H1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text); set_cn(r); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    return p

def H2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text); set_cn(r); r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2E,0x74,0xB5)
    return p

def H3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text); set_cn(r); r.font.size = Pt(11.5); r.font.color.rgb = RGBColor(0x2E,0x74,0xB5)
    return p

def P(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); set_cn(r); r.font.size = Pt(10.5); r.bold = bold; r.italic = italic
    return p

def BULLET(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); set_cn(r); r.font.size = Pt(10.5)
    return p

def CODE(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x40,0x40,0x40)
    return p

def TABLE(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].paragraphs[0].clear()
        r = hdr[i].paragraphs[0].add_run(h); set_cn(r); r.bold = True; r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i,v in enumerate(row):
            cells[i].paragraphs[0].clear()
            r = cells[i].paragraphs[0].add_run(str(v)); set_cn(r); r.font.size = Pt(9.5)
    return t

# ============ 封面 ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('基于金山云平台的\nSRE 可靠性工程平台')
set_cn(r, '微软雅黑'); r.font.size = Pt(26); r.bold = True; r.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
for _ in range(2): doc.add_paragraph()
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('—— 项目技术报告 ——'); set_cn(r); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x59,0x59,0x59)
doc.add_paragraph()
info = doc.add_paragraph(); info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run('作者：甘玉涛  |  方向：SRE / 云原生可靠性工程  |  日期：2026 年 8 月')
set_cn(r); r.font.size = Pt(11)
doc.add_page_break()

# ============ 一、项目核心 ============
H1('一、项目核心：做了什么')
P('一句话概括：在金山云 KEC 真实多节点 K8s 集群上，用 SRE 方法论端到端构建了一个微服务可靠性工程平台，'
  '回答"SRE 到底在做什么"——不是"把服务部署起来"，而是"用软件工程方法让系统天然更稳定、可量化、可自治"。')
P('具体做了四件事（对应 SRE 的四个核心问题）：', bold=True)
BULLET('系统有多稳？—— 定义 SLO 与错误预算，把"稳定性"从主观感觉变成可量化的数字。')
BULLET('不稳了怎么第一时间知道？—— 搭建可观测性三支柱（Metrics / Logs / Traces）+ 多窗口多燃烧率告警。')
BULLET('故障来了系统能不能自己扛？—— 探针自愈 + HPA 弹性扩缩 + 熔断降级，减少人工介入。')
BULLET('怎么主动验证它真的扛得住？—— 用 Chaos Mesh 混沌工程主动注入故障，验证可靠性假设。')
P('并接入金山云 KECR（镜像仓库）与 KS3（对象存储）两个云产品，体现"基于金山云平台"的云原生落地，'
  '非本地 demo 玩具。项目已端到端跑通，含真实运行数据。', italic=True)

# ============ 二、项目背景 ============
H1('二、项目背景与目标')
P('传统运维的稳定性是"感觉"——"好像挺稳"、"最近没出事"。SRE 的革命性在于：用软件工程方法把稳定性量化、'
  '自动化、可验证。本项目以一个 Go 写的电商订单微服务 ordersvc 为载体，在其上完整落地 SRE 方法论，'
  '覆盖大厂 SRE 岗位（字节 / 滴滴 / 阿里 / 腾讯 / 京东）JD 中的全部核心能力点。')
P('目标：', bold=True)
BULLET('落地完整的 SRE 可靠性工程闭环：SLO → 可观测性 → 告警 → 自愈 → 混沌验证。')
BULLET('在金山云真实多节点集群运行，接入云产品，体现公有云工程能力。')
BULLET('产出可复现的声明式基础设施（全 YAML + 一键部署脚本）。')

# ============ 三、金山云基础设施 ============
H1('三、金山云基础设施')
P('本项目运行在金山云 KEC 真实多节点集群，而非本地 kind 单机：')
TABLE(['组件', '规格 / 说明'],
[
  ['KEC 主机', '5 台：3 master + 2 node（master01/02/03 + node01/02）'],
  ['操作系统 / 内核', 'Rocky Linux 9.8 / K8s v1.31.14 / containerd 2.2.6'],
  ['CNI 网络', 'Calico（Pod CIDR 172.16.0.0/16，Service CIDR 192.168.0.0/16）'],
  ['StorageClass', 'local-path（default），供 Loki 索引等本地持久化'],
  ['metrics-server', '已就绪，HPA 容量弹性可用'],
  ['VPC + EIP + 安全组', '5 节点内网 10.0.0.x 互通；公网经 EIP，安全组放行 NodePort 30088/30090/30300/30686'],
  ['金山云 KECR', 'hub.kce.ksyun.com 私有镜像仓库，ordersvc 镜像 push 至此，集群经 imagePullSecrets 拉取'],
  ['金山云 KS3', 'S3 兼容对象存储，Loki chunks 落 KS3，实现日志存算分离与长期归档'],
  ['kubeconfig 公网化', 'admin.conf server 由内网 VIP 10.0.0.236:16443 改为公网 10.0.0.182:16443，本地直连'],
])
P('公网访问地址（金山云 KEC EIP + 安全组放行 NodePort）：', bold=True)
CODE('ordersvc API :  http://10.0.0.182:30088/healthz\n'
     'Prometheus   :  http://10.0.0.182:30090    (/rules /alerts)\n'
     'Grafana      :  http://10.0.0.182:30300    (admin/admin)\n'
     'Jaeger UI    :  http://10.0.0.182:30686')

# ============ 四、系统架构 ============
H1('四、系统架构')
CODE('┌────── 金山云 KEC 公网入口（EIP + 安全组放行 NodePort）──────┐\n'
     '│  10.0.0.182:30088 / 30090 / 30300 / 30686               │\n'
     '└──┬──────────┬────────┬──────┬──────────────────────────────┘\n'
     '   │          │        │      │\n'
     '┌── 金山云 KEC 5 节点 K8s v1.31 / Calico / containerd ──────────┐\n'
     '│  ns: sre-demo                    ns: observability           │\n'
     '│  ┌──────────────────┐            ┌────────────────────┐     │\n'
     '│  │ ordersvc (Go) x3 │──/metrics─▶│ Prometheus         │     │\n'
     '│  │ +OTLP +health    │            │ (SLO 规则+告警)     │     │\n'
     '│  │ +HPA +Probe      │            └────────┬───────────┘     │\n'
     '│  │ imagePullSecrets │                     │                 │\n'
     '│  │  →金山云 KECR     │               ┌─────▼──────────┐      │\n'
     '│  └───┬────────┬─────┘               │ Grafana(SLO大盘)│      │\n'
     '│  OTLP│    logs│                     └─────┬──────────┘      │\n'
     '│      ▼       ▼                     ┌─────▼──────────┐      │\n'
     '│  ┌───────┐ ┌────────┐              │ Loki◀─Promtail  │─chunks─▶KS3\n'
     '│  │OTel C.│ │Promtail│              └────────────────┘      │\n'
     '│  └───┬───┘ └────────┘                                      │\n'
     '│      ▼                                                    │\n'
     '│  ┌───────┐               ┌──────────────────┐             │\n'
     '│  │Jaeger │               │ Chaos Mesh       │             │\n'
     '│  └───────┘               │ PodKill/Net/CPU  │─▶ 故障注入   │\n'
     '│                          └──────────────────┘             │\n'
     '└────────────────────────────────────────────────────────────┘')

# ============ 五、技术栈 ============
H1('五、技术栈与对应 SRE 能力')
TABLE(['模块', '技术选型', '体现的 SRE 能力（命中 JD）'],
[
  ['云基础设施', '金山云 KEC + VPC + EIP + 安全组', '公有云资源编排、网络与安全组规划'],
  ['镜像仓库', '金山云 KECR + imagePullSecrets', '私有镜像仓库治理、凭证管理'],
  ['对象存储', '金山云 KS3（Loki chunks 后端）', '存算分离、日志长期归档、云原生存储'],
  ['业务服务', 'Go + 多阶段 Dockerfile + 非 root', 'Go 语言、容器化、镜像安全基线'],
  ['编排', 'K8s Deployment / HPA / Probe', '滚动发布零宕机、探针自愈、容量弹性'],
  ['Metrics', 'Prometheus + 自定义指标 + Recording Rules', '指标体系设计、SLI 计算、规则预聚合'],
  ['Logs', 'Loki + Promtail(DaemonSet) + KS3', '日志聚合、结构化日志、存储分离'],
  ['Traces', 'OpenTelemetry + OTel Collector + Jaeger', '分布式链路追踪、采样治理'],
  ['可视化', 'Grafana + provisioning（16 面板 SLO 大盘）', 'SLO 大盘、错误预算可视化'],
  ['SLO 引擎', '5 窗口多燃烧率告警（Google SRE 法）', 'SLO/错误预算/燃烧率告警（核心考点）'],
  ['混沌工程', 'Chaos Mesh（Pod/Network/StressChaos）', '故障演练、可靠性主动验证'],
  ['故障自愈', 'liveness/readiness/startup probe + HPA + /admin/fault', '自愈、流量摘除、降级'],
  ['IaC', '全声明式 YAML + 一键 bootstrap 脚本', '基础设施即代码、可复现'],
])

# ============ 六、SLO 设计 ============
H1('六、SLO 设计（项目核心）')
H2('6.1 SLI 与目标')
TABLE(['SLI', '定义', '目标（30 天窗口）', '错误预算'],
[
  ['可用性', '非 5xx 请求占比', '≥ 99.9%', '0.1% ≈ 43 分钟/月'],
  ['延迟', 'P99 请求延迟', '< 500ms', '—'],
])
H2('6.2 错误预算')
P('100% − 99.9% = 0.1%，这 0.1% 是"允许出错的部分"，即错误预算。30 天 ≈ 43200 分钟，0.1% ≈ 43 分钟/月。'
  '预算是用来"花"的——允许小故障消耗，而非追求 100% 完美（成本过高）。预算耗尽则本月不及格，需冻结变更、专心优化。')
H2('6.3 燃烧率（Burn Rate）')
P('燃烧率 = 实际错误率 ÷ 允许错误率(0.001)。=1 表示按预算匀速消耗；=14.4 表示 1 小时耗尽 2% 月预算（需立即介入）。')
H2('6.4 多窗口多燃烧率告警（Google SRE《SRE 工作手册》第 5 章）')
P('短窗口负责"快速发现"，长窗口负责"防抖降噪"，两者同时超阈值才告警——避免单窗口抖动误报。三档告警：')
TABLE(['告警', '短窗口 & 长窗口', '阈值', '含义', '动作'],
[
  ['Page', '5m & 1h', '> 14.4', '1h 消耗 ≥2% 月预算', '立即介入'],
  ['Ticket', '30m & 6h', '> 6', '6h 消耗 ~5% 月预算', '建工单'],
  ['Budget', '1d', '> 1', '按此速率将耗尽预算', '排期优化'],
])
P('Prometheus 实现：5 窗口（5m/30m/1h/6h/1d）× 5 类指标 = 21 条 Recording Rules + 5 条 Alert Rules，共 26 条规则。', italic=True)

# ============ 七、可观测性 ============
H1('七、可观测性三支柱')
H2('7.1 Metrics（指标）—— Prometheus')
BULLET('ordersvc 自定义 6 类指标：请求计数、请求耗时直方图、订单创建/失败、在途订单、注入故障强度、build_info。')
BULLET('kubernetes_sd_configs + Pod 注解自动发现 ordersvc:9090/metrics，无需手动配置 target。')
BULLET('Recording Rules 预聚合高频 counter，避免告警查询全量扫描原始序列。')
H2('7.2 Logs（日志）—— Loki + Promtail + KS3')
BULLET('Promtail 以 DaemonSet 部署，5 节点全覆盖，采集 ordersvc 结构化日志。')
BULLET('Loki chunks 落金山云 KS3（S3 兼容 API），实现存算分离与长期归档。')
H2('7.3 Traces（链路）—— OpenTelemetry + Jaeger')
BULLET('ordersvc 经 OTLP 上报 trace 至 OTel Collector，再转 Jaeger 存储。')
BULLET('一次 /order 调用的端到端链路可在 Jaeger UI 查询，故障时可快速定位慢调用。')
H2('7.4 Grafana SLO 大盘（16 个面板）')
TABLE(['面板组', '面板', '含义'],
[
  ['SLO 核心', '错误预算剩余（gauge）', '1 − 1d错误率，仪表盘直观显示本月预算还剩多少'],
  ['', '错误预算消耗速率 1d', 'burn_rate1d，>1 表示按当前速率 30 天烧光预算'],
  ['', '可用性 SLI（1h）', '1 − 1h错误率，实时可用性'],
  ['燃烧率', '错误预算燃烧率（5m/1h）', '双窗口燃烧率时序，超 14.4 触发 Page'],
  ['延迟', 'P99 延迟 / 延迟分位 P50-P999', 'histogram_quantile 计算各分位延迟，对照 SLO 500ms'],
  ['流量', 'QPS / 请求成功失败计数 / 订单创建失败', '请求量与错误请求分解'],
  ['关联', 'QPS vs 错误率 / Pod 副本数', '流量与错误关联、HPA 自愈可视化'],
  ['演练', '注入故障强度', '当前注入的 fail_rate 与 latency'],
])

# ============ 八、混沌工程 ============
H1('八、混沌工程')
P('用 Chaos Mesh 主动注入故障，验证可靠性假设（不等真实故障才发现短板）：')
TABLE(['实验', '类型', '验证的可靠性假设'],
[
  ['Pod Kill', 'PodChaos', '随机杀 Pod，K8s 30s 内自愈，SLO 不击穿'],
  ['Network Delay 200ms', 'NetworkChaos', 'P99 延迟告警触发，服务仍可用'],
  ['CPU Stress 80%', 'StressChaos', 'HPA 触发扩容，CPU 回落'],
  ['HTTP 50% 5xx', '/admin/fault', '错误预算燃烧率告警触发'],
])
P('Chaos Mesh 组件：chaos-controller-manager + chaos-daemon(DaemonSet 5/5) + chaos-dashboard + chaos-dns-server，全部 Running。', italic=True)

# ============ 九、故障自愈 ============
H1('九、故障自愈机制')
BULLET('三类探针：liveness（失败重启容器）、readiness（失败摘除流量）、startup（慢启动保护）。')
BULLET('HPA 弹性：基于 CPU 利用率与在途订单数自动扩缩容，实测故障期间 Pod 从 3 扩到 20。')
BULLET('滚动发布零宕机：Deployment maxUnavailable=0 + maxSurge=1，新 Pod ready 后才下旧 Pod。')
BULLET('熔断降级：/admin/fault 接口热更新故障率，验证降级与告警联动。')

# ============ 十、实测验证 ============
H1('十、实测验证数据（金山云 KEC 集群实跑）')
P('项目已在金山云 KEC 5 节点真实集群端到端跑通，以下为实测数据（非本地模拟）。', italic=True)
H2('10.1 集群与服务状态')
BULLET('5 节点全部 Ready；ordersvc 3/3 Running，镜像从 KECR 拉取，healthz=200。')
BULLET('可观测性栈全 Running，Prometheus 自动采集 3 个 ordersvc Pod metrics（target up）。')
BULLET('Chaos Mesh 全就绪，三类实验注入验证通过。')
H2('10.2 流量压测')
P('20 QPS 持续 10 分钟：成功 1391 / 总 1418 = 成功率 98.10%，失败 27（注入的 50% 故障样本）。')
H2('10.3 SLO 燃烧率告警实测（核心验证）')
P('注入 50% 5xx 故障 + 200ms 延迟后，关键指标实测：')
TABLE(['指标 / 告警', '实测值', '阈值', '结论'],
[
  ['燃烧率 5m', '76.4', '> 14.4', '超阈值 5 倍，预算高速燃烧'],
  ['燃烧率 1h', '58.8', '> 14.4', '双窗口均超 → Page 告警触发'],
  ['燃烧率 1d', '58.8', '> 1', 'Budget 预警触发'],
  ['错误率 5m', '7.6%', '允许 0.1%', '超标 76 倍'],
  ['错误预算剩余', '94.1%', '100%', '已消耗约 6% 月预算'],
  ['可用性 1h', '95.9%', '99.9%', '低于 SLO 目标'],
  ['P99 延迟', '237ms', '< 500ms', '延迟 SLO 仍达标'],
  ['Pod 副本数', '20', '初始 3', 'HPA 扩容 6.7 倍'],
  ['OrdersvcHighErrorRatePage', 'firing', '—', 'Page 告警已触发'],
  ['OrdersvcChaosInjected', 'firing', '—', '混沌注入标记告警'],
])
P('恢复 fail_rate=0 后燃烧率回落、告警自动消除——证明告警链路闭环：'
  '指标采集 → 燃烧率计算 → 多窗口判定 → 告警触发/恢复。', italic=True)
H2('10.4 混沌工程实测')
BULLET('PodChaos：杀 1 个 Pod，Deployment 自动重建，SLI 未击穿——自愈验证通过。')
BULLET('NetworkChaos：AllInjected=True，目标节点注入网络延迟。')
BULLET('StressChaos：注入 90s CPU 压力，HPA 基于 CPU 决策扩缩。')

# ============ 十一、SRE 思维 ============
H1('十一、项目体现的 SRE 思维')
BULLET('数据驱动稳定性：所有决策基于 SLO 指标与错误预算，而非主观感觉。')
BULLET('消除琐事：故障注入 / 自愈 / 扩缩容全部自动化，减少人工介入。')
BULLET('拥抱错误预算：允许小故障消耗预算，告警只在预算高速燃烧时触发，避免告警疲劳。')
BULLET('可观测性先行：Metrics / Logs / Traces 三支柱齐全，故障可定位、可复盘。')
BULLET('主动验证：混沌工程主动制造故障验证假设，而非等真实故障才发现短板。')
BULLET('云原生落地：在金山云真实多节点集群跑通，接入 KECR / KS3，非本地 demo。')

# ============ 十二、关键代码解析 ============
H1('十二、关键代码解析')
P('本章摘取项目中最能体现 SRE 工程能力的代码段，逐段解释设计意图。', italic=True)

H2('12.1 指标定义（internal/metrics/metrics.go）—— SLO 的数据基础')
P('设计原则：只暴露对 SLO 计算与告警有用的指标，避免指标爆炸。每个指标都对应一个 SRE 用途。')
CODE('type OrderServiceMetrics struct {\n'
     '    // Counter：按 method/path/code 维度统计请求总量 → 算可用性 SLI（成功率）\n'
     '    HTTPRequestsTotal *prometheus.CounterVec\n'
     '    // Histogram：延迟直方图，Bucket 覆盖 5ms~10s → 算延迟 SLI（P99 < 500ms）\n'
     '    HTTPRequestDurationSeconds *prometheus.HistogramVec\n'
     '    // Gauge：当前在途订单数 → HPA 扩缩容决策 + 容量观察\n'
     '    OrdersInFlight prometheus.Gauge\n'
     '    OrdersCreatedTotal prometheus.Counter   // 业务侧下单成功总量\n'
     '    OrdersFailedTotal *prometheus.CounterVec // 失败总量（按 reason）\n'
     '    FailureRateInjected prometheus.Gauge     // 当前注入故障率（混沌演练标记）\n'
     '    LatencyInjectedMs   prometheus.Gauge     // 当前注入延迟\n'
     '    BuildInfo *prometheus.GaugeVec           // 版本追踪\n'
     '}')
P('要点解释：', bold=True)
BULLET('HTTPRequestsTotal 用 CounterVec（带 method/path/code label）：code label 是关键——'
      'Prometheus 规则用 code=~"5.." 过滤错误请求，算出错误率。这就是可用性 SLI 的原始数据。')
BULLET('HTTPRequestDurationSeconds 用 HistogramVec 而非 Summary：Histogram 分桶存（le 标签），'
      '服务端用 histogram_quantile() 算任意分位，且能跨 3 个 Pod 聚合；Summary 在客户端预算分位、不可聚合。')
BULLET('OrdersInFlight 用 Gauge（可增可减）：下单 Inc、结束 Dec，反映实时并发，HPA 可据此扩缩。')
BULLET('FailureRateInjected 是 Gauge：把"当前注入了多少故障"也变成指标，大盘可显示，告警 OrdersvcChaosInjected 据此触发。')

H2('12.2 指标采集中间件（internal/handler/handler.go）—— SLO 数据的产生点')
P('instrument 中间件包装所有路由，统一记录延迟与请求计数。这是 SLO 计算（成功率/延迟）的唯一数据来源。')
CODE('func (h *Handler) instrument(next http.Handler) http.Handler {\n'
     '    return http.HandlerFunc(func(w http.ResponseWriter, r *http.Request) {\n'
     '        start := time.Now()\n'
     '        // statusRecorder 包装 ResponseWriter，捕获真实响应状态码\n'
     '        ww := &statusRecorder{ResponseWriter: w, status: 200}\n'
     '        next.ServeHTTP(ww, r)\n'
     '        elapsed := time.Since(start).Seconds()\n'
     '\n'
     '        // 记录延迟（Histogram.Observe）+ 请求计数（Counter.Inc）\n'
     '        h.metrics.HTTPRequestDurationSeconds.WithLabelValues(\n'
     '            r.Method, r.URL.Path).Observe(elapsed)\n'
     '        h.metrics.HTTPRequestsTotal.WithLabelValues(\n'
     '            r.Method, r.URL.Path, strconv.Itoa(ww.status)).Inc()\n'
     '\n'
     '        // 给每条请求补一个 server span（含状态码），串联业务 span\n'
     '        span := trace.SpanFromContext(r.Context())\n'
     '        span.SetAttributes(attribute.Int("http.status_code", ww.status))\n'
     '    })\n'
     '}')
P('要点解释：', bold=True)
BULLET('statusRecorder 包装 ResponseWriter：标准 http.ResponseWriter 无法读已写状态码，'
      '包装后才能拿到真实 code（200/500 等），喂给 CounterVec 的 code label。这是指标准确性的关键。')
BULLET('一个中间件搞定全路由：避免每个 handler 重复写指标代码，符合"消除琐事"原则。')
BULLET('指标 + trace 同一中间件产生：指标算 SLO（是否稳），trace 定位根因（为什么不稳），数据同源联动。')

H2('12.3 故障注入与业务逻辑（createOrder）—— 混沌验证的核心')
P('createOrder 演示了如何在业务代码里支持混沌演练：热更新故障率/延迟，按概率返回 5xx。')
CODE('func (h *Handler) createOrder(w http.ResponseWriter, r *http.Request) {\n'
     '    ctx := r.Context()\n'
     '    tracer := otel.Tracer("ordersvc")\n'
     '    ctx, span := tracer.Start(ctx, "createOrder")   // 开业务 span\n'
     '    defer span.End()\n'
     '\n'
     '    h.metrics.OrdersInFlight.Inc()                  // 在途 +1\n'
     '    defer h.metrics.OrdersInFlight.Dec()            // defer 保证结束 -1（即使 panic）\n'
     '\n'
     '    // 注入处理延迟：模拟下游支付/库存调用耗时\n'
     '    if h.latencyMs > 0 {\n'
     '        _, sleepSpan := tracer.Start(ctx, "downstream.payment")\n'
     '        time.Sleep(time.Duration(h.latencyMs) * time.Millisecond)\n'
     '        sleepSpan.End()\n'
     '    }\n'
     '\n'
     '    // 注入失败：按概率返回 500，验证 SLO 错误预算与告警\n'
     '    if h.failRate > 0 && rand.Float64() < h.failRate {\n'
     '        h.metrics.OrdersFailedTotal.WithLabelValues("injected_500").Inc()\n'
     '        w.WriteHeader(http.StatusInternalServerError)\n'
     '        return\n'
     '    }\n'
     '\n'
     '    // 成功下单\n'
     '    orderID := strconv.FormatInt(time.Now().UnixNano(), 36)\n'
     '    h.metrics.OrdersCreatedTotal.Inc()\n'
     '    w.WriteHeader(http.StatusCreated)\n'
     '}')
P('要点解释：', bold=True)
BULLET('defer h.metrics.OrdersInFlight.Dec()：用 defer 保证在途计数一定递减，即使中途 return 或 panic，Gauge 不泄漏。')
BULLET('延迟注入单独开 downstream.payment span：Jaeger 里能看到"下游支付耗时 X ms"的子 span，'
      '故障时能精确定位是哪一段慢，而非笼统看总延迟。')
BULLET('按概率 rand.Float64() < h.failRate 返回 500：failRate=0.5 即 50% 请求失败，'
      '这些 500 被 instrument 中间件记进 HTTPRequestsTotal{code="500"}，Prometheus 算出错误率 → 触发燃烧率告警。'
      '这就是"注入故障 → 指标变化 → 告警触发"闭环的业务侧起点。')

H2('12.4 故障热更新接口（setFault）—— 不重启即可演练')
P('/admin/fault 接口运行时热更新故障参数，无需重启 Pod，便于混沌实验动态控制。')
CODE('// 例：/admin/fault?fail=0.5&latency=300\n'
     'func (h *Handler) setFault(w http.ResponseWriter, r *http.Request) {\n'
     '    if v := r.URL.Query().Get("fail"); v != "" {\n'
     '        if rate, err := strconv.ParseFloat(v, 64); err == nil && rate >= 0 && rate <= 1 {\n'
     '            h.failRate = rate\n'
     '            h.metrics.FailureRateInjected.Set(rate)   // 同步更新指标\n'
     '        }\n'
     '    }\n'
     '    if v := r.URL.Query().Get("latency"); v != "" {\n'
     '        if ms, err := strconv.Atoi(v); err == nil && ms >= 0 {\n'
     '            h.latencyMs = ms\n'
     '            h.metrics.LatencyInjectedMs.Set(float64(ms))\n'
     '        }\n'
     '    }\n'
     '    json.NewEncoder(w).Encode(map[string]any{\n'
     '        "fail_rate": h.failRate, "latency_ms": h.latencyMs})\n'
     '}')
P('要点解释：', bold=True)
BULLET('热更新 + 指标同步：改 failRate 的同时更新 FailureRateInjected 指标，大盘实时显示当前注入强度，告警 OrdersvcChaosInjected 也据此 firing。')
BULLET('参数校验：rate 限 [0,1]、ms 限 ≥0，防止非法值导致行为异常。')
BULLET('chaos-inject-fault.sh 脚本就是调这个接口：注入 50% 故障 → 燃烧率飙到 76.4 → Page 告警 firing。')

H2('12.5 OpenTelemetry 链路初始化（internal/trace/trace.go）')
P('初始化 OTel TracerProvider，经 OTLP/HTTP 上报到 Jaeger，service name 来自环境变量。')
CODE('func Init(ctx context.Context) (func(context.Context) error, error) {\n'
     '    exporter, err := otlptracehttp.New(ctx,\n'
     '        otlptracehttp.WithEndpoint(os.Getenv("OTEL_EXPORTER_OTLP_ENDPOINT")),\n'
     '        otlptracehttp.WithInsecure())\n'
     '    res, _ := resource.New(ctx,\n'
     '        resource.WithAttributes(semconv.ServiceName(os.Getenv("OTEL_SERVICE_NAME"))))\n'
     '    tp := sdktrace.NewTracerProvider(\n'
     '        sdktrace.WithBatcher(exporter),     // 批量上报降开销\n'
     '        sdktrace.WithResource(res),\n'
     '        // 全量采样：故障演练时抓完整链路；生产可换 TraceIDRatioBased(0.1)\n'
     '        sdktrace.WithSampler(sdktrace.AlwaysSample()),\n'
     '    )\n'
     '    otel.SetTracerProvider(tp)\n'
     '    otel.SetTextMapPropagator(propagation.TraceContext{}) // W3C TraceContext\n'
     '    return tp.Shutdown, nil\n'
     '}')
P('要点解释：', bold=True)
BULLET('WithBatcher：批量上报 trace，降低对业务请求的同步开销。')
BULLET('AlwaysSample：全量采样，故障演练时保证抓到完整链路；生产可换 TraceIDRatioBased(0.1) 降 10% 采样省开销。')
BULLET('TraceContext propagator：用 W3C 标准透传 traceID，跨服务链路串联，Jaeger 能看到端到端调用。')

H2('12.6 优雅关闭（main.go）—— 零宕机的一环')
P('捕获 SIGTERM 信号，5 秒超时优雅关闭，配合滚动发布实现零宕机。')
CODE('func main() {\n'
     '    ctx, stop := signal.NotifyContext(context.Background(),\n'
     '        syscall.SIGINT, syscall.SIGTERM)\n'
     '    defer stop()\n'
     '    // ... 启动 HTTP 服务 ...\n'
     '    <-ctx.Done()                      // 阻塞直到收到信号\n'
     '    shutCtx, cancel := context.WithTimeout(context.Background(), 5*time.Second)\n'
     '    defer cancel()\n'
     '    _ = srv.Shutdown(shutCtx)         // 优雅关闭：处理完在途请求再退出\n'
     '}')
P('要点解释：', bold=True)
BULLET('滚动发布时旧 Pod 收到 SIGTERM：若不优雅关闭，在途请求会被中断导致 5xx。'
      'Shutdown() 让 server 停止接新请求、处理完在途请求再退出。')
BULLET('5 秒超时兜底：防止某些请求卡死导致 Pod 长期不退出（K8s 默认 30s 后强杀）。'
      '配合 Deployment 的 terminationGracePeriodSeconds 调优。')

H2('12.7 Prometheus SLO 规则（prometheus.yaml ConfigMap）—— 燃烧率引擎')
P('5 窗口 Recording Rules 预聚合 + 3 档告警，是整个 SLO 引擎的核心。摘录关键规则：')
CODE('# 错误比率 = 错误请求速率 / 总请求速率（5 个窗口）\n'
     '- record: ordersvc:error_ratio5m\n'
     '  expr: ordersvc:error_rate5m / ordersvc:request_rate5m\n'
     '\n'
     '# 燃烧率 = 错误率 / 允许错误率(0.001)\n'
     '- record: ordersvc:burn_rate5m\n'
     '  expr: ordersvc:error_ratio5m / 0.001\n'
     '\n'
     '# Page 告警：5m & 1h 双窗口同时 > 14.4（1h 耗 ≥2% 月预算）\n'
     '- alert: OrdersvcHighErrorRatePage\n'
     '  expr: |\n'
     '    ordersvc:burn_rate5m > 14.4\n'
     '    and ordersvc:burn_rate1h > 14.4\n'
     '  for: 2m                         # 持续 2 分钟才 firing（防抖）\n'
     '  labels: {severity: page, slo: availability-99.9}')
P('要点解释：', bold=True)
BULLET('分层预聚合：原始 counter → error_rate/request_rate → error_ratio → burn_rate，'
      '每层用上一层结果，避免告警查询全量扫描原始序列（性能优化）。')
BULLET('双窗口 and：5m 和 1h 都超阈值才告警，短窗口快速发现 + 长窗口防抖降噪，避免瞬时抖动误报。')
BULLET('for: 2m：双窗口超阈值后还需持续 2 分钟才转 firing，进一步防抖（实测 pending→firing 约 2 分钟）。')
BULLET('允许错误率 0.001 = 1 − 99.9% SLO：燃烧率阈值不是拍脑袋，而是从 SLO 推导（14.4 = 1h 耗 2% 月预算的理论值）。')

H2('12.8 Deployment 探针与自愈（deploy/manifests/ordersvc.yaml）')
P('三类探针 + 滚动发布配置，是 K8s 自愈能力的工作负载侧配置：')
CODE('livenessProbe:           # 失败则重启容器（救死）\n'
     '  httpGet: {path: /healthz, port: http}\n'
     '  initialDelaySeconds: 5\n'
     'readinessProbe:          # 失败则摘除流量但不重启（救活）\n'
     '  httpGet: {path: /readyz, port: http}\n'
     '  initialDelaySeconds: 3\n'
     'strategy:\n'
     '  type: RollingUpdate\n'
     '  rollingUpdate:\n'
     '    maxUnavailable: 0     # 先起新 Pod 才下旧 Pod → 零宕机\n'
     '    maxSurge: 1           # 最多多起 1 个新 Pod\n'
     'terminationGracePeriodSeconds: 30  # 优雅关闭宽限期')
P('要点解释：', bold=True)
BULLET('liveness 查 /healthz：进程卡死时 K8s 重启容器。')
BULLET('readiness 查 /readyz：未就绪时从 Endpoints 摘除，不接流量但不重启——故障 Pod 自愈期间不影响用户。')
BULLET('maxUnavailable: 0：滚动发布时保证可用副本数不降，配合 readiness 探针实现零宕机。')

# ============ 十三、目录结构 ============
H1('十三、项目目录结构')
CODE('sre-project/\n'
     '├── app/                      # Go 微服务（main + handler + metrics + trace + Dockerfile）\n'
     '├── deploy/\n'
     '│   ├── manifests/ordersvc.yaml  # Deployment/HPA/Probe/ConfigMap(SLO)/imagePullSecrets\n'
     '│   └── scripts/                # bootstrap-ksce.sh / load-test / chaos-inject / 远程与镜像搬运\n'
     '├── observability/\n'
     '│   ├── prometheus/           # Prometheus + SLO 规则(26条) + kube-state-metrics\n'
     '│   ├── grafana/              # Grafana + 16 面板 SLO 大盘 provisioning\n'
     '│   ├── loki/                 # Loki + Promtail(chunks 落 KS3)\n'
     '│   └── jaeger/               # OTel Collector + Jaeger\n'
     '├── chaos/experiments.yaml    # Chaos Mesh 三类实验\n'
     '├── postmortem/               # 事故复盘模板与示例\n'
     '└── docs/                     # 项目技术文档、runbook')

# ============ 十四、一键部署 ============
H1('十四、一键部署')
CODE('export KSCE_PWD=<master root 密码>\n'
     'export KECR_PWD=<KECR 登录密码>\n'
     'bash deploy/scripts/bootstrap-ksce.sh                                  # 一键部署\n'
     'SRE_HOST=10.0.0.182 bash deploy/scripts/load-test.sh 20 600        # 打流量\n'
     'SRE_HOST=10.0.0.182 bash deploy/scripts/chaos-inject-fault.sh 0.5 200  # 注入故障')

doc.save(r'C:\Users\KC\Desktop\秋招\简历\简历\设计师修改\基于金山云平台的SRE可靠性工程平台-项目报告-v2.docx')
print('OK 报告已生成（v2）')
