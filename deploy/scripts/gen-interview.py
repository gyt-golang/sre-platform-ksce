#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《基于金山云平台的 SRE 可靠性工程平台》面试面经 Word 版。
按技术栈分章，每章含：技术点说明 + 高频面试问题 + 参考答案（结合本项目）。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_cn(run, name='微软雅黑'):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def H1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text); set_cn(r); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    return p
def H2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text); set_cn(r); r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2E,0x74,0xB5)
    return p
def H3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text); set_cn(r); r.font.size = Pt(11.5); r.font.color.rgb = RGBColor(0x2E,0x74,0xB5)
    return p
def P(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); set_cn(r); r.font.size = Pt(10.5); r.bold = bold; r.italic = italic
    return p
def BULLET(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); set_cn(r); r.font.size = Pt(10.5)
    return p
def CODE(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x40,0x40,0x40)
    return p
def Q(text):
    """面试问题"""
    p = doc.add_paragraph()
    r = p.add_run('Q：'); set_cn(r); r.bold = True; r.font.color.rgb = RGBColor(0xC0,0x39,0x2B); r.font.size = Pt(11)
    r = p.add_run(text); set_cn(r); r.bold = True; r.font.size = Pt(11)
    return p
def A(text):
    """参考答案"""
    p = doc.add_paragraph()
    r = p.add_run('A：'); set_cn(r); r.bold = True; r.font.color.rgb = RGBColor(0x1E,0x7E,0x34); r.font.size = Pt(10.5)
    r = p.add_run(text); set_cn(r); r.font.size = Pt(10.5)
    return p
def TABLE(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].paragraphs[0].clear()
        r = hdr[i].paragraphs[0].add_run(h); set_cn(r); r.bold = True; r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i,v in enumerate(row):
            cells[i].paragraphs[0].clear()
            r = cells[i].paragraphs[0].add_run(str(v)); set_cn(r); r.font.size = Pt(9.5)
    return t

# ============ 封面 ============
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('基于金山云平台的\nSRE 可靠性工程平台'); set_cn(r,'微软雅黑'); r.font.size = Pt(26); r.bold = True; r.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
for _ in range(2): doc.add_paragraph()
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('—— 面试面经（技术栈高频问答）——'); set_cn(r); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x59,0x59,0x59)
doc.add_paragraph()
info = doc.add_paragraph(); info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run('甘玉涛  |  SRE / 云原生方向  |  2026 年 8 月'); set_cn(r); r.font.size = Pt(11)
doc.add_page_break()

# ============ 开篇 ============
H1('一、面试讲解主线（3 分钟讲清）')
P('按这个顺序讲，逻辑最顺：')
BULLET('背景：传统运维稳定性靠感觉，我想用 SRE 方法论量化它。')
BULLET('做了什么：在金山云 KEC 5 节点真实集群上，围绕 ordersvc 微服务落地 SRE 四件事——SLO 定义、可观测性三支柱、错误预算告警、混沌工程验证。')
BULLET('核心亮点：多窗口多燃烧率告警（Google SRE 法），实测注入 50% 故障 → 燃烧率 76.4 → Page 告警 firing，HPA 自动扩容 3→20。')
BULLET('云原生落地：接入金山云 KECR 镜像仓库 + KS3 对象存储，非本地 demo。')
P('下面按技术栈分章，每章列高频问题与参考答案（结合本项目实际）。', italic=True)

# ============ 二、技术栈总览 ============
H1('二、技术栈总览')
TABLE(['类别', '技术', '本项目用途'],
[
  ['云平台', '金山云 KEC / VPC / EIP / 安全组', '5 节点集群、公网入口、NodePort 暴露'],
  ['云产品', 'KECR / KS3', '私有镜像仓库 / Loki 对象存储'],
  ['语言', 'Go', 'ordersvc 微服务'],
  ['容器', 'Docker / containerd', '多阶段构建、非 root 镜像'],
  ['编排', 'Kubernetes v1.31', 'Deployment / HPA / Probe / Service'],
  ['Metrics', 'Prometheus', '指标采集 + Recording Rules + Alert'],
  ['Logs', 'Loki + Promtail', '日志聚合 + KS3 存储'],
  ['Traces', 'OpenTelemetry + Jaeger', '分布式链路追踪'],
  ['可视化', 'Grafana', '16 面板 SLO 大盘'],
  ['混沌', 'Chaos Mesh', 'Pod / Network / Stress Chaos'],
  ['IaC', 'Helm / kubectl / YAML', 'Chaos Mesh 部署、声明式资源'],
])

# ============ 三、SLO 与错误预算 ============
H1('三、SLO 与错误预算（核心考点）')
H2('技术点')
BULLET('SLI：可用性（非 5xx 占比）≥ 99.9%；延迟 P99 < 500ms。')
BULLET('错误预算 = 1 − SLO = 0.1% ≈ 43 分钟/月。')
BULLET('燃烧率 = 实际错误率 / 允许错误率(0.001)。')
BULLET('多窗口多燃烧率：5m&1h(Page) / 30m&6h(Ticket) / 1d(Budget)。')
H2('高频问答')
Q('什么是 SLO？和 SLA、SLI 有什么区别？')
A('SLI 是测量指标（如成功率、P99 延迟）；SLO 是 SLI 的目标值（如成功率 ≥ 99.9%）；'
  'SLA 是对外承诺的合同（违约有赔偿，通常比 SLO 更宽松）。本项目定义可用性 SLI = 非 5xx 占比，SLO = 99.9%。')
Q('什么是错误预算？有什么用？')
A('错误预算 = 1 − SLO = 0.1%/月 ≈ 43 分钟。它是"允许出错的额度"。用途：① 预算没用完时可放心发版；'
  '② 预算耗尽则冻结变更、专心优化；③ 用燃烧率衡量预算消耗速度，决定告警级别。本项目注入 50% 故障后，'
  '1d 燃烧率达 58.8，错误预算剩余从 100% 跌到 94.1%。')
Q('什么是燃烧率？为什么用多窗口？')
A('燃烧率 = 实际错误率 / 允许错误率。=1 匀速消耗，=14.4 表示 1h 烧 2% 月预算。'
  '多窗口（如 5m & 1h 双窗口同时超阈值才告警）：短窗口快速发现，长窗口防抖降噪，避免单窗口抖动误报。'
  '本项目 Page 告警 = 5m>14.4 AND 1h>14.4，实测 5m=76.4、1h=58.8 同时触发。')
Q('为什么选 99.9% 而不是 99% 或 99.99%？')
A('SLO 不是越高越好——越高成本越大（99.9%→99.99% 成本通常翻 10 倍）。99.9% 是大多数在线服务的甜点区，'
  '错误预算 43 分钟/月足够覆盖正常发版与小故障。选型应基于用户期望与成本权衡，而非盲目追高。')
Q('预算耗尽怎么办？')
A('冻结非关键变更，专心修问题；用错误预算"负债"机制——若本月超支，下月预算扣减。'
  '本项目 Budget 告警（1d burn>1）即预警"照这速度 30 天烧光"，提前排期优化。')

# ============ 四、Prometheus ============
H1('四、Prometheus（Metrics）')
H2('技术点')
BULLET('架构：Pull 模型，Prometheus 主动抓 target；TSDB 本地存储；Recording Rules 预聚合。')
BULLET('服务发现：kubernetes_sd_configs + Pod 注解（prometheus.io/scrape）自动发现。')
BULLET('指标类型：Counter（递增）/ Gauge（瞬时）/ Histogram（分桶）/ Summary。')
BULLET('PromQL：rate() 算速率、histogram_quantile() 算分位。')
H2('高频问答')
Q('Prometheus 为什么用 Pull 而不是 Push？')
A('Pull 优势：① 主动控频，避免被大量 target 打爆；② target 不用关心 Prom 是否在线；'
  '③ 可通过 /metrics 端点人工调试；④ Prom 宕机不影响业务。Push 适合短生命周期任务（如 CronJob），'
  '此时用 Pushgateway 中转。本项目 ordersvc 暴露 /metrics，Prom 经 Pod 注解自动 Pull。')
Q('Counter 和 Gauge 区别？Histogram 和 Summary 区别？')
A('Counter 只增不减（如请求总数），用 rate() 算 QPS；Gauge 可增可减（如在途订单数）。'
  'Histogram 把数据分桶存（le 标签），服务端用 histogram_quantile() 算任意分位，跨实例可聚合；'
  'Summary 在客户端预算分位，不可聚合。本项目延迟用 Histogram，便于跨 Pod 聚合算 P99。')
Q('什么是 Recording Rules？为什么用？')
A('把高频计算的 PromQL 预聚合存成新指标，避免告警查询时全量扫描原始序列。'
  '本项目把 ordersvc_http_requests_total 预聚合成 5 窗口 × 5 类共 21 条规则（如 ordersvc:burn_rate5m），'
  '告警直接查预聚合指标，性能好且语义清晰。')
Q('Prometheus 服务发现怎么做？')
A('本项目用 kubernetes_sd_configs role=pod + relabel_configs：'
  'keep 带 prometheus.io/scrape=true 注解的 Pod，replace 注解里的 path/port。'
  'ordersvc 加注解后自动被发现，无需手动配 target。还部署了 kube-state-metrics 暴露 K8s 对象指标。')
Q('Prometheus 怎么算 P99 延迟？')
A('histogram_quantile(0.99, sum by (le, path) (rate(ordersvc_http_request_duration_seconds_bucket[5m])))。'
  '先 rate 算各桶速率，sum by le 聚合，再 quantile 算分位。本项目实测 P99=237ms，低于 500ms SLO。')

# ============ 五、Kubernetes ============
H1('五、Kubernetes 编排')
H2('技术点')
BULLET('Deployment：滚动发布（maxUnavailable=0/maxSurge=1）零宕机。')
BULLET('三类探针：liveness / readiness / startup。')
BULLET('HPA：基于 CPU / 自定义指标自动扩缩。')
BULLET('Service：NodePort 对外暴露。')
H2('高频问答')
Q('liveness、readiness、startup 探针区别？')
A('liveness：失败重启容器（救死）；readiness：失败摘除流量但不重启（救活，未就绪不接请求）；'
  'startup：慢启动保护，启动完成前不跑 liveness/readiness（避免启动慢被误杀重启循环）。'
  '本项目 ordersvc 三类都配，readiness 保证故障 Pod 不接流量。')
Q('滚动发布怎么做到零宕机？')
A('Deployment 默认 RollingUpdate：maxUnavailable=0（先起新 Pod 才下旧 Pod）、maxSurge=1（最多多起 1 个）。'
  '配合 readinessProbe，新 Pod ready 才进入 Endpoints，旧 Pod 才终止。'
  '还可加 terminationGracePeriodSeconds 让旧 Pod 处理完在途请求。')
Q('HPA 怎么工作的？本项目扩到几个？')
A('HPA 周期查询指标（CPU 利用率/自定义指标），超 target 利用率按公式算期望副本数 = 当前 × (当前指标/target)。'
  '本项目 HPA 基于 CPU + 在途订单数，故障注入 CPU 压力后实测从 3 副本扩到 20。需 metrics-server 提供 CPU 指标。')
Q('Service NodePort 和 LoadBalancer 区别？')
A('NodePort 在所有节点开固定端口（30000-32767），任一节点 IP:端口可访问，简单但暴露所有节点。'
  'LoadBalancer 自动调云 LB（需 cloud-controller-manager），公网只暴露 LB IP。'
  '本项目用 NodePort（自建集群无 CCM），金山云安全组放行 30088 等端口。')
Q('Pod 一直 CrashLoopBackOff 怎么排查？')
A('kubectl logs 看应用日志；kubectl describe pod 看 Events（OOMKilled/ImagePullBackOff/探针失败）；'
  'kubectl get events 看集群事件。本项目曾遇 chaos-daemon CrashLoopBackOff——日志报连 docker.sock 超时，'
  '根因是设了 runtime=containerd 但没设 socketPath=/run/containerd/containerd.sock，改后恢复。')

# ============ 六、可观测性 Logs/Traces ============
H1('六、可观测性：Logs 与 Traces')
H2('技术点')
BULLET('Loki：日志聚合，只索引 label 不索引全文，轻量；Promtail DaemonSet 采集。')
BULLET('KS3：Loki chunks 落对象存储，存算分离。')
BULLET('OpenTelemetry：厂商中立 trace 标准；OTel Collector 接 OTLP 转 Jaeger。')
H2('高频问答')
Q('Loki 和 ELK 区别？为什么选 Loki？')
A('ELK（Elasticsearch）全文索引，功能强但重、成本高；Loki 只索引 label（像 Prometheus 的日志版），'
  '日志原文存对象存储，轻量便宜，与 Grafana/Prometheus 同一套 label 体系联动。'
  '本项目日志量不大、追求与指标联动，选 Loki + KS3 存算分离。')
Q('Loki 怎么做存储分离？')
A('Loki 把 chunks（日志块）存金山云 KS3（S3 兼容 API），索引/chunks 元数据存本地 PV（local-path SC）。'
  'common.storage 配 s3 + endpoint/bucket/AK/SK。存算分离让 Loki 无状态化、可水平扩展、日志长期归档。')
Q('OpenTelemetry 是什么？为什么用它？')
A('OTel 是 CNCF 的可观测性数据采集标准（统一 trace/metrics/logs），厂商中立，避免绑定 Jaeger/Zipkin/Prom。'
  'ordersvc 用 OTel SDK 经 OTLP 协议上报 trace → OTel Collector → Jaeger 存储。换后端只改 Collector 配置，业务无感。')
Q('Metrics / Logs / Traces 三支柱怎么联动定位故障？')
A('① Prometheus 告警发现"错误率升"；② Grafana 看是哪个 path/实例；'
  '③ Loki 按 label 过滤该实例结构化日志看错误详情；④ Jaeger 按 traceID 看慢调用链路定位根因。'
  '三支柱用同一套 label（namespace/pod/path）关联，本项目大盘已集成。')

# ============ 七、混沌工程 ============
H1('七、混沌工程（Chaos Mesh）')
H2('技术点')
BULLET('Chaos Mesh：CNCF 混沌平台，CRD 定义实验（PodChaos/NetworkChaos/StressChaos）。')
BULLET('chaos-daemon DaemonSet 覆盖所有节点，执行实际注入。')
BULLET('一次性实验 vs Schedule 周期实验。')
H2('高频问答')
Q('混沌工程是什么？和价值在哪？')
A('主动、可控地注入故障（杀 Pod、网络延迟、CPU 压力），验证系统可靠性假设，提前发现短板。'
  '价值：不等真实故障才暴露问题；验证自愈/告警/容灾是否真生效；建立系统韧性信心。'
  '本项目用 Chaos Mesh 注入 PodKill/NetDelay/CPUStress，验证 K8s 自愈与 HPA 扩容。')
Q('Chaos Mesh 架构？chaos-daemon 干啥？')
A('chaos-controller-manager 解析 CRD 实验并调度；chaos-daemon 以 DaemonSet 跑在每个节点，'
  '实际执行注入（调容器运行时、iptables、tc 等）；chaos-dashboard 可视化。'
  '本项目 daemon 5/5 覆盖 5 节点，NetworkChaos/StressChaos 经 daemon 注入。')
Q('注入网络延迟实验，服务怎么自愈？')
A('NetworkChaos 注入 200ms 延迟 → Prometheus P99 延迟指标上升 → OrdersvcHighLatencyP99 告警触发。'
  '服务本身因有超时与重试仍可用，SLO 不一定击穿。实验到期自动恢复（AllRecovered=True）。')
Q('Chaos Mesh 2.7 踩过什么坑？')
A('① spec.scheduler 字段已移除（周期调度改用独立 Schedule CRD），单 Chaos 资源做一次性实验；'
  '② chaosDaemon.runtime=containerd 必须同时设 socketPath=/run/containerd/containerd.sock，否则连 docker.sock 超时；'
  '③ chaos-daemon 镜像在 ghcr.io，国内需配镜像加速或节点间搬运。')

# ============ 八、容器与镜像 ============
H1('八、容器、镜像与 KECR')
H2('技术点')
BULLET('多阶段构建：builder 阶段编译，运行阶段只 COPY 二进制，镜像小、攻击面小。')
BULLET('非 root 运行（USER 65532），符合安全基线。')
BULLET('KECR 私有仓库 + imagePullSecrets。')
H2('高频问答')
Q('为什么要多阶段构建？')
A('把构建环境（含 Go 工具链、源码）与运行环境分离。builder 阶段编译出静态二进制，运行阶段用 alpine 只 COPY 二进制。'
  '镜像从几百 MB 降到十几 MB，减少攻击面（无编译器/源码），拉取快、安全。本项目 CGO_ENABLED=0 静态编译。')
Q('镜像为什么用非 root？')
A('容器与宿主共享内核，root 容器一旦逃逸即获宿主 root。非 root（USER 65532:nobody）符合容器安全基线，'
  '即使被攻破权限也有限。本项目 ordersvc 以 65532 运行。')
Q('imagePullSecrets 怎么工作？')
A('创建 docker-registry Secret（含仓库地址/用户名/密码），在 Deployment spec.imagePullSecrets 引用。'
  'kubelet 拉私有镜像时用该 Secret 认证。本项目 ordersvc 镜像在金山云 KECR，集群经 regcred Secret 拉取。')
Q('containerd 镜像加速怎么配？')
A('containerd 1.7+ 用 config_path 模式：/etc/containerd/certs.d/<registry>/hosts.toml 配 mirror host。'
  '本项目 5 节点配了 docker.io / ghcr.io / registry.k8s.io 三个加速（daocloud），解决国内拉取超时。'
  '改 hosts.toml 动态生效，无需重启 containerd。')

# ============ 九、Go 微服务 ============
H1('九、Go 微服务设计')
H2('技术点')
BULLET('暴露 /metrics（Prometheus 格式）、/healthz、/readyz、/admin/fault。')
BULLET('Histogram 指标（延迟分桶）、Counter（请求计数）。')
BULLET('OTLP 上报 trace。')
H2('高频问答')
Q('为什么 ordersvc 用 Go？')
A('大厂 SRE/云原生首选（字节/滴滴 JD 明确）：编译型、并发原语好（goroutine）、镜像小、启动快；'
  'K8s/Prometheus/Chaos Mesh 生态同语言，二次开发友好。')
Q('应用怎么暴露指标给 Prometheus？')
A('用 prometheus/client_golang 注册 Collector：Counter（http_requests_total 按 code/path label）、'
  'Histogram（request_duration_seconds_bucket 按 le 分桶）。/metrics 端点输出 Prometheus 文本格式。'
  '加 prometheus.io/scrape 注解让 Prom 自动发现。')
Q('/admin/fault 接口干啥用？')
A('热更新应用内故障率与延迟（fail_rate / latency_ms），不重启即可模拟 5xx 与慢响应。'
  '用于验证错误预算燃烧率告警链路——注入 50% 故障后 burn_rate5m 飙到 76.4，Page 告警 firing。')

# ============ 十、金山云平台 ============
H1('十、金山云平台与云原生落地')
H2('技术点')
BULLET('KEC：5 节点云主机；VPC 内网互通；EIP + 安全组公网入口。')
BULLET('KECR：私有镜像仓库；KS3：对象存储（S3 兼容）。')
BULLET('kubeconfig 公网化：改 server 为公网 IP + insecure-skip-tls-verify。')
H2('高频问答')
Q('为什么用金山云 KEC 而不是本地 kind？')
A('真实多节点集群才能验证：① 跨节点调度与自愈；② DaemonSet（Promtail/chaos-daemon）多节点覆盖；'
  '③ 真实网络（Calico 跨节点通信）；④ 接入云产品（KECR/KS3）。kind 单机无法体现这些，且面试官更看重真实云环境落地。')
Q('kubeconfig 怎么从内网改公网访问？')
A('原 admin.conf 指向内网 VIP 10.0.0.236:16443（haproxy 高可用入口）。改成公网 master01 IP 10.0.0.182:16443。'
  '因 apiserver 证书 SAN 不含公网 IP，加 insecure-skip-tls-verify: true 跳过证书校验（demo 用法，生产应重新签发含公网 IP 的证书）。')
Q('KS3 在项目里怎么用？为什么？')
A('Loki chunks 落 KS3（S3 兼容 API），实现日志存算分离。原因：① 日志长期归档不占本地盘；'
  '② Loki 无状态化可水平扩展；③ 对象存储便宜。体现云原生"存算分离"架构。')
Q('安全组要放行哪些端口？')
A('NodePort 30088（ordersvc）/ 30090（Prometheus）/ 30300（Grafana）/ 30686（Jaeger）。'
  '安全组是金山云 VPC 的虚拟防火墙，入向规则控制公网可访问端口。')

# ============ 十一、故障排查与实战 ============
H1('十一、故障排查与实战经验')
H2('高频问答')
Q('可观测性 Pod 全部 ImagePullBackOff，怎么排查？')
A('kubectl describe pod 看 Events——若 Failed to pull image + i/o timeout，是镜像拉不下来。'
  '本项目根因：5 节点中只有 master01 配了 docker.io 镜像加速，其余 4 节点直连 docker.io 超时。'
  '解决：给所有节点配 /etc/containerd/certs.d/docker.io/hosts.toml 加速器。')
Q('chaos-daemon 一直 Terminating 删除不了？')
A('无 finalizer 但容器停止 hang（privileged 进程不响应），旧 Pod 占 slot 导致 DaemonSet 不建新 Pod。'
  'kubectl delete pod --force --grace-period=0 强制从 etcd 删除，DaemonSet 立即重建。本项目多次用此恢复。')
Q('注入故障后告警一直 pending 不 firing？')
A('Page 告警有 for:2m 条件——双窗口燃烧率超阈值后需持续 2 分钟才转 firing（防抖）。'
  'pending 是正常中间态。本项目实测注入 50% 故障约 2 分钟后转 firing。')
Q('项目最大的收获是什么？')
A('① 把 SRE 方法论从书本落到真实云环境，可量化可验证；'
  '② 踩了大量真实坑（镜像加速/RBAC/Chaos Mesh 兼容性），积累实战排障经验；'
  '③ 理解了"数据驱动稳定性"——所有决策基于指标与错误预算，而非感觉。')

# ============ 十二、高频追问预设 ============
H1('十二、高频追问预设')
Q('这个项目最难的地方在哪？')
A('不是写代码，而是把 SRE 方法论端到端跑通并验证。难点：① 多窗口燃烧率告警的规则设计与 PromQL 实现；'
  '② 在真实多节点集群解决镜像拉取（docker.io/ghcr.io/registry.k8s.io 三套加速 + 节点间搬运）；'
  '③ Chaos Mesh 2.7 与 containerd 的兼容性（socket path / scheduler 字段变更）。')
Q('如果让你扩展，你会做什么？')
A('① Prometheus 持久化用 Thanos + KS3 长期存储；② 引入金丝雀发布（Argo Rollouts）做变更管理；'
  '③ 用 K6 做更专业的压测；④ Loki Ruler 日志告警接 Alertmanager 统一出口；'
  '⑤ 告警驱动自愈，告警触发自动 restart/scale，形成检测→告警→自愈闭环。')
Q('错误预算告警和普通阈值告警（如错误率>5%）有什么本质区别？')
A('普通阈值告警凭经验拍阈值，与 SLO 脱节；错误预算告警从 SLO 推导，阈值有理论依据（14.4 = 1h 耗 2% 月预算），'
  '且多窗口天然分级（Page/Ticket/Budget 对应不同介入力度），避免告警疲劳。这是 SRE 与传统运维监控的本质差异。')
Q('SLO 定义之后怎么持续运营？')
A('三步闭环：① slo-spec.yaml 作为单一事实源声明 SLO 目标/SLI/告警阈值，'
  'recording/alert rules 由 spec 派生，避免手写 26 条规则易错；② slo-report.py 周期查 Prometheus '
  '生成 SLO 达成报告（SLI/错误预算剩余/燃烧率），主动回顾而非被动等告警；③ 不达标时按报告调整 SLO 目标或排期优化。')
Q('toil（人工劳动）怎么量化和管理？')
A('toil-log.py 记录每次手动干预（任务/耗时/可自动化），toil-report.py 聚合算成本（元/分钟）排自动化优先级。'
  '本项目回填 5 条 toil 共 84min、可自动化 64min（76%）、可回收 160 元，ks3-integration-debug 45min 列 P0。'
  'SRE 目标 toil < 50% 工时，超出即驱动自动化。')
Q('postmortem 怎么保证不只是写文档？')
A('① schema.json 定义必填章节（概要/影响/根因/行动项），validate-postmortem.py 校验；'
  '② 行动项强制 ≥1 个 DONE/IN-PROGRESS，禁止全 TODO 堆积，证明跟踪到闭环；'
  '③ blameless 无指责复盘。本次混沌演练 postmortem 的"打 chaos label 抑制告警"行动项已 DONE，'
  '落地为 Alertmanager inhibit_rules。')
Q('了解监控配置即代码（Monitoring as Code）吗？')
A('Mixin 理念：SLI 定义、告警规则、Grafana dashboard、runbook 是同一 SLO 的四个视图，'
  '应打包版本化、一起部署一起 review（Jsonnet 实现，如 sre-monitoring-as-code）。'
  '本项目 slo-spec.yaml 是单一事实源，rules/dashboard/runbook 由 spec 派生对齐；'
  'demo 手动对齐，生产可用自动生成工具彻底消除四者漂移。')

doc.save(r'C:\Users\KC\Desktop\秋招\简历\简历\设计师修改\基于金山云平台的SRE可靠性工程平台-面经.docx')
print('OK 面经已生成')
