#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《SRE 项目面试深挖面经》Word 版。

专攻"项目级深挖问答"——面试官听完项目介绍后的连环追问：
整体设计/技术权衡/踩坑复盘/性能容量/故障场景/延伸挑战。
与 gen-interview.py（基础技术点面经）互补。

用法：python deploy/scripts/gen-interview-deep.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def set_cn(run, name='微软雅黑'):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def TITLE(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); set_cn(r); r.font.size = Pt(20); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def SUB(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); set_cn(r); r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def H1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text); set_cn(r); r.font.size = Pt(15); r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def H2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text); set_cn(r); r.font.size = Pt(12.5); r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)


def Q(text):
    p = doc.add_paragraph()
    r = p.add_run('Q：' + text); set_cn(r); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)


def A(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.left_indent = Pt(12); pf.space_after = Pt(6)
    r = p.add_run('A：' + text); set_cn(r); r.font.size = Pt(10.5)


# ============ 封面 ============
TITLE('SRE 项目面试深挖面经')
SUB('项目级连环追问 · 设计权衡 / 踩坑复盘 / 故障场景 / 延伸挑战')
SUB('基于金山云 KEC 平台的 SRE 可靠性工程平台 · 2026-08')
doc.add_paragraph()

# ============ 一、整体设计与灵魂三问 ============
H1('一、整体设计与灵魂三问')
A('面试官开场多半会问"这个项目解决什么问题、为什么这么设计"。回答要有一条主线，不能散讲工具。')

Q('用一句话说清这个项目解决了什么问题？')
A('传统运维是"出事我去手动救"，本项目用软件工程方法让系统"自己知道自己有多稳、不稳会自动告警自愈、救完会复盘改进"。'
  '核心是把稳定性从拍脑袋的玄学，变成用错误预算驱动的、可量化可自治的工程问题，全套跑在金山云 5 节点真实集群上。')

Q('为什么不用现成的 SRE 平台（如 Datadog/Grafana Cloud），要自己搭？')
A('一是面试要体现"能从零搭起来"的工程能力，不是只会用 SaaS；二是自搭能深入理解每层原理（Pull 模型、燃烧率阈值推导、存算分离），'
  '用 SaaS 是黑盒；三是成本可控，金山云 KS3 对象存储比 SaaS 便宜。生产里确实会用托管服务，但懂底层才能在出问题时定位。')

Q('这个项目里你最得意的一个设计是什么？')
A('错误预算驱动的多窗口燃烧率告警。它一举解决了三个问题：①告警阈值有理论依据（14.4=1h 烧 2% 月预算，不是拍脑袋）；'
  '②多窗口天然分级（Page/Ticket/Budget 对应不同介入力度）；③允许小故障消耗预算不打扰，只在高速燃烧时告警，避免告警疲劳。'
  '这是 SRE 和传统运维监控的本质区别。')

Q('如果让你给这个项目打分，哪里做得好、哪里是短板？')
A('好的：SLO 全链路闭环跑通且有实测数据、云产品真实接入（KECR/KS3）、运营闭环（toil 量化/postmortem 校验）是差异化亮点。'
  '短板：①单节点 Loki/Prometheus 没做高可用，生产需集群化；②没做告警驱动的自动修复（自愈仅 K8s 原生）；'
  '③没接入真实通知渠道压力测试。这些是已知的"下一步"，面试主动说出来比被问出来好。')

# ============ 二、技术选型与权衡 ============
H1('二、技术选型与权衡')
A('面试官爱问"为什么选 A 不选 B"，考的是你是否理解每个选择的代价。')

Q('为什么 SLO 目标定 99.9% 而不是 99% 或 99.99%？')
A('SLO 不是越高越好。99% 太松（每月允许 7.2h 故障，用户体验差）；99.99% 太贵（成本通常比 99.9% 翻 10 倍，需多活+容灾）。'
  '99.9%（每月 43.2 分钟翻车额度）是大多数在线服务的甜点区，用户基本无感且成本可控。SLO 应和业务价值、成本三者平衡，不是追求四个九。')

Q('Prometheus 为什么用 Pull 而不是 Push？Pull 不会漏数据吗？')
A('Pull 的好处：①主动控制采集节奏，易于做容量规划；②目标挂了 Prometheus 能立刻发现（scrape 失败即告警），Push 模式目标挂了你反而不知道；'
  '③无状态、易调试（curl 即可看指标）。漏数据问题用 Push Gateway 补（短任务场景）。本项目 ordersvc 常驻服务，Pull 完全够用。')

Q('Loki 为什么不像 ELK 那样给日志建全文索引？不怕查得慢吗？')
A('这是 Loki 的核心设计哲学：只索引 label（namespace/pod/level），不索引日志正文。'
  '好处：索引体积小、写入快、存储省（chunks 落 KS3 对象存储极便宜）；适合 K8s 场景（label 天然有结构）。'
  '代价：全文检索慢。但日志排查 90% 是"某 pod 某时段某 level"的过滤查询，label 索引够用。要全文检索的场景才用 ES。这是用空间换场景适配。')

Q('为什么日志存储用对象存储（KS3）而不是块存储/PVC？')
A('存算分离的考量：①容量近乎无限按需扩，不受单盘限制；②Pod/节点挂了日志不丢（对象存储独立于集群）；'
  '③Loki 变无状态可水平扩展，本地只存 wal/cache；④对象存储比块存储便宜约 10 倍，适合海量日志长期归档。'
  '块存储适合需要低延迟随机读的数据库，日志是"写多读少量大"，对象存储正合适。')

Q('Chaos Mesh 和自己写故障注入脚本，为什么要两个都做？')
A('应用层脚本（/admin/fault）注入业务逻辑故障（返回 5xx、加延迟），验证的是 SLO 告警链路；'
  'Chaos Mesh 注入基础设施故障（杀 Pod、网络延迟、CPU 压力），验证的是 K8s 自愈和韧性。'
  '两者层次不同、验证目标不同。Chaos Mesh 还能做稳态假设、实验编排，是体系化的混沌工程，脚本只是轻量补充。')

Q('金丝雀发布为什么能比滚动发布更安全？AnalysisTemplate 自动回滚的判断依据是什么？')
A('滚动发布是"按批次替换 Pod"，新版本一旦 Rolling 出去就开始接真实流量，发现问题只能人工 rollback，'
  '中间窗口的故障影响所有命中新 Pod 的用户。金丝雀是"按流量比例切流"（20%→40%→60%→100%），'
  '每一步 pause 跑 AnalysisTemplate 查 Prometheus SLO 指标，超阈值自动 abort 并回滚到上一个 stable ReplicaSet，'
  '故障被限制在 20% 流量内且秒级回退。'
  'AnalysisTemplate 的判断依据是 SLO 指标而非 Pod Ready：Ready 只说明进程活着，不代表"对用户没伤害"；'
  '查错误率（5xx 占比 ≤5%）和 P99 延迟（≤800ms）才是"变更是否真的健康"。'
  '阈值比 SLO 宽（SLO 错误率 0.1%，金丝雀 5%），因为新版本刚切流样本少，5% 是"明显有问题"的硬阈值，'
  '触发即回滚，避免样本噪声误杀正常发布。本质是把"发布是否健康"从人肉看大盘变成指标驱动的自动决策。')

# ============ 三、踩坑与故障复盘 ============
H1('三、踩坑与故障复盘')
A('面试官特别爱问踩坑——考的是你真的动手做过，还是纸上谈兵。每个坑都要能讲清"现象→定位→根因→解决"。')

Q('Loki 接 KS3 你踩过什么坑？怎么定位的？')
A('踩了四个：①S3 配置字段名错——Loki 3.x 是 bucketnames/s3forcepathstyle（不是老式 bucket_name），Pod CrashLoop，看启动日志 "field not found in type aws.S3Config" 定位；'
  '②Promtail 没配 ServiceAccount RBAC，日志报 "pods is forbidden"，发现不了 Pod 不采集——但 Loki 不报错，series 永远空，最难查；'
  '③Promtail 缺 __path__ relabel，报 "no path for target" 不知道读哪个日志文件；'
  '④集群 local-path provisioner 动态建 PV 超时 120s，改 emptyDir（反正索引已上传 KS3）。'
  '教训：可观测组件"静默失败"最坑，必须查 series/labels 确认数据真的流通，不能只看 Pod Running。')

Q('Alertmanager 你之前没部署，怎么发现的？')
A('对照代码和文档发现的——postmortem 里写"Alertmanager 触发告警"，但 prometheus.yaml 的 alertmanagers.targets 是空的，集群里也没有 alertmanager pod。'
  '规则会 firing 但没有路由、通知、抑制，链路断在最后一环。这是"代码与文档脱节"的典型——文档先按"已做"写了，实际没落地。'
  '教训：文档里每个"已做"都要回代码核实，否则面试官一查 GitHub 就露馅。')

Q('容器停止卡在 Terminating，你怎么处理的？')
A('容器进程不响应 SIGTERM（或存储卸载慢）导致超时挂起。用 kubectl delete pod --force --grace-period=0 强制删除。'
  '根因要具体分析：若是进程不响应，加 preStop hook 给优雅退出时间；若是挂载点卡住，查 kubelet/容器运行时。'
  'demo 里直接 force delete，生产里要找根因避免依赖强制删除。')

Q('混沌演练时告警疯狂 firing，怎么避免打扰？')
A('给演练注入的故障打 chaos=true label，Alertmanager 配 inhibit_rules：chaos=true 时抑制 Page 告警，'
  '演练故障不打扰真实 oncall。同时 info 级的 OrdersvcChaosInjected 路由到 null-receiver 仅 UI 可见。'
  '这是告警治理的典型场景——告警要能区分"真实事故"和"演练噪音"。')

Q('金丝雀发布验证自动回滚时，坏版本发出去却没回滚，怎么定位的？')
A('现象：发了会 50% 返回 500 的 v3-bad 版本，金丝雀一路推进到 100% 成了 stable，AnalysisRun 全是 Successful，没触发回滚。'
  '定位分两步：①先 exec 进 v3-bad pod 直接 curl /order，确认 10 次里有 8 次 500——buggy 标志确实生效，问题不在镜像；'
  '②再去 Prometheus 手动跑 AnalysisTemplate 里的错误率 PromQL，发现结果恒为 0。'
  '根因：指标 ordersvc_http_requests_total 的状态码 label 名是 code（见 metrics.go 定义），'
  '但 AnalysisTemplate 的 PromQL 写的是 status=~"5.."，label 名对不上 → 匹配到空 series → '
  '分子为空 → 叠加查询末尾的 or vector(0) 兜底返回 0 → 错误率永远 0 → 永远 ≤5% 阈值 → Analysis 永远 success。'
  '修复：把 PromQL 的 status 改成 code。改完重测，错误率立刻读到 0.33，AnalysisRun 转 Failed，v3-bad RS 缩到 0，自动回滚到 v2。'
  '教训：①or vector(0) 这种"无流量兜底"是双刃剑——它让无流量时不报 NaN，但也把"label 写错导致查不到数据"的真实故障掩盖成 0；'
  '②canary 健康分析的 PromQL 必须和指标定义的 label 名逐字核对，最好先在 Prometheus 手动验证查询有值再放进 AnalysisTemplate；'
  '③"分析一直成功"本身就是可疑信号——金丝雀期样本少，全 success 要警惕是不是查询根本没匹配到数据。')

# ============ 四、性能与容量 ============
H1('四、性能与容量')
A('SRE 面试必问容量——"你这能扛多大量、瓶颈在哪、怎么扩"。')

Q('这套系统单节点能扛多少 QPS？瓶颈在哪？')
A('ordersvc 是轻量 Go 服务，单实例数千 QPS 没问题（实测打 600 并发稳定）。瓶颈不在业务，在可观测性栈：'
  '①Prometheus 单实例 ingest 能力约 1M samples/s，超了要分片（分 namespace 起多实例 + Thanos 聚合）；'
  '②Loki 单实例 ingest 受 ingester 限流（本项目就触发了 429 限流），扩要多 ingester + 分发器；'
  '③KS3 对象存储本身近无限，瓶颈在 Loki 的 ingest 速率。生产扩容先扩可观测性栈，不是业务。')

Q('Prometheus 数据存哪？长期数据怎么解决？')
A('本项目 Prometheus 用 emptyDir（demo 规模够）。生产不能这样——Pod 重启数据全丢。'
  '方案：①短期（15天）存本地 PVC；②长期用 Thanos，把 Prometheus 数据定期上传到对象存储（金山云 KS3），'
  'Thanos Query 跨实例聚合查询，实现全局视图+长期存储+高可用。本质是复用 KS3 做 Prom 的持久层。')

Q('Loki 日志量大了怎么办？怎么控制成本？')
A('①采样子集 pod 全量日志、其余采样；②用 pipeline_stages 丢弃调试级日志（只留 warn+）；'
  '③retention_period 按需设（热数据 7 天 KS3，冷数据归档更便宜的存储类）；④多 ingester 分担写入；'
  '⑤label 基数控制（别用 user_id 这种高基数当 label，否则索引爆炸）。成本控制是 SRE 运营能力的体现。')

Q('Grafana 大盘查询慢，怎么排查？')
A('①看是不是查了过大时间范围（30天全量）——缩短范围或用 recording rule 预聚合；'
  '②看是不是高基数 label 导致 series 爆炸——label_values 查 series 数；'
  '③看 PromQL 有无低效写法（如 sum(rate) 嵌套过深、没用 recording rule）；'
  '④Prometheus 本身慢——看 query duration 指标，考虑加 Thanos 或分片。')

# ============ 五、故障场景与应急 ============
H1('五、故障场景与应急')
A('给你一个故障场景让你处置——考的是 oncall 思维和排障逻辑，不是背答案。')

Q('凌晨告警 ordersvc 错误率飙升到 30%，你怎么处置？')
A('按 oncall SOP：①先看告警带的 runbook_url/dashboard_url，30 秒内掌握上下文（哪个 SLO、错误预算消耗多少）；'
  '②看 Grafana 确认是真实故障还是误报（对比历史基线）；③看是不是最近有变更（发版/配置改动）——80% 故障是变更引起；'
  '④看日志（Loki 按 namespace+pod 过滤）和链路（Jaeger 看 slow span）定位根因；'
  '⑤若是代码 bug，rollback；若是容量问题，扩容；若是依赖故障，切降级。'
  '⑥处置中持续更新事故状态，事后 postmortem。核心：先止血再治本，MTTR 优先。')

Q('订单服务突然延迟 P99 飙到 2 秒，但错误率没涨，怎么查？')
A('错误率没涨说明请求还在成功，是变慢不是失败。排查方向：①Jaeger 看链路，定位慢在哪个 span（DB？下游调用？）；'
  '②看资源水位（CPU/内存/网络）是否打满——可能是流量涨了但 HPA 没跟上；③看依赖（DB 连接池、缓存命中率）；'
  '④看 GC（Go 服务的 STW）。延迟问题比错误问题更难查，因为"还能用"容易被忽视，但已在烧延迟 SLO 的预算。')

Q('整个可用区挂了（金山云某机房故障），你的系统怎么办？')
A('本项目是单集群单机房，扛不住可用区级故障——这是已知短板。生产方案：①多集群多可用区部署，流量跨机房负载均衡；'
  '②Prometheus/Alertmanager 多副本跨机房；③Loki 数据在 KS3（对象存储本身跨机房冗余），可跨集群重建；'
  '④DNS/全局负载均衡做机房级故障转移。本项目先验证方法论，多活是下一步。')

Q('KS3 挂了（对象存储故障），Loki 还能工作吗？')
A('不能正常写入新 chunks（ingester 会报错堆积），但已写入的历史日志不丢（假设 KS3 只是暂时不可达）。'
  '缓解：①ingester 有 wal 本地缓冲，短期断连能扛；②Prometheus 告警 KS3 不可达，oncall 介入；'
  '③极端情况切到本地 filesystem 降级（牺牲存算分离换可用）。生产里对象存储可用性通常很高（多副本），'
  '但 SLO 要把对象存储可用性算进依赖链。')

# ============ 六、延伸挑战与开放题 ============
H1('六、延伸挑战与开放题')
A('面试官最后常问"如果让你继续做，你会做什么"——考技术视野和优先级判断。')

Q('如果让你把这个项目做到生产级，你会怎么改？')
A('按优先级：①可观测性栈高可用——Prometheus/Loki/Alertmanager 多副本，消除单点；'
  '②Thanos 接 KS3 做 Prom 长期存储+全局查询；③告警真正接入飞书/钉钉并做压测（通知风暴场景）；'
  '④告警驱动自愈——Page 告警触发自动 restart/scale（现在自愈仅 K8s 原生）；'
  '⑤多集群多可用区；⑥接入真实业务流量验证容量。优先级依据：先消除单点（影响最大），再补长存储，最后做高级自愈。')

Q('SRE 和运维的本质区别是什么？')
A('工具可能一样，本质是思维和方法：①运维是反应式救火（比 MTTR），SRE 是预防式工程（比事故频率和预算消耗）；'
  '②运维把稳定当状态（"没告警就是稳"），SRE 把稳定当可计算预算；③运维靠经验，SRE 靠数据驱动闭环；'
  '④运维靠人，SRE 靠自动化消除 toil；⑤运维是手工作坊，SRE 是软件工程（一切 as code）。'
  '一句话：从经验驱动的反应式运维，升级成数据驱动的工程式可靠性管理。')

Q('你觉得做 SRE 最重要的能力是什么？')
A('不是会多少工具，是"系统思维+数据驱动+工程化"。系统思维：能从全局看稳定性（指标→告警→自愈→复盘是闭环不是散点）；'
  '数据驱动：每个决策有数据依据（SLO/错误预算/toil 量化），不靠感觉；'
  '工程化：把稳定性做成代码和自动化（slo-spec、Alertmanager 配置、自愈 playbook），可复用可演进。'
  '工具会过时，这三点是底层能力，也是 SRE 比运维值钱的地方。')

Q('如果 SLO 一直达标，是不是就不用管了？')
A('不是。SLO 持续达标可能意味着 SLO 定低了（用户实际体验比目标好很多），可以适当上调 SLO 释放更多变更空间，'
  '或降低成本（缩容）。也可能掩盖了"靠堆人力达成"的问题——要看 toil 占比，若 toil 高说明稳定性是靠人换的，不可持续。'
  'SLO 是要持续运营调优的，不是定完就完。这也是为什么有 slo-report.py 周期回顾。')

doc.save(r'C:\Users\KC\Desktop\秋招\简历\简历\设计师修改\SRE项目面试深挖面经.docx')
print('OK 深挖面经已生成：SRE项目面试深挖面经.docx')
